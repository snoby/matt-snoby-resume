from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/snoby/Dropbox/Playground/matt-snoby-resume/Matthew_Snoby_Resume_ATS_2026_reworked.docx"


def set_cell_margins(section):
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)


def set_paragraph_border(paragraph, color="D9D9D9", size=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def ensure_style(document, name, style_type, base_name=None):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, style_type)
    if base_name:
        style.base_style = styles[base_name]
    return style


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    name_style = ensure_style(document, "Resume Name", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    name_style.font.name = "Calibri"
    name_style.font.size = Pt(18)
    name_style.font.bold = True
    name_style.font.color.rgb = RGBColor(31, 54, 95)
    name_style.paragraph_format.space_after = Pt(1)

    title_style = ensure_style(document, "Resume Title", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(12.5)
    title_style.font.bold = False
    title_style.font.color.rgb = RGBColor(68, 68, 68)
    title_style.paragraph_format.space_after = Pt(2)

    contact_style = ensure_style(document, "Resume Contact", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    contact_style.font.name = "Calibri"
    contact_style.font.size = Pt(9.5)
    contact_style.font.color.rgb = RGBColor(68, 68, 68)
    contact_style.paragraph_format.space_after = Pt(7)

    section_style = ensure_style(document, "Section Heading", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    section_style.font.name = "Calibri"
    section_style.font.size = Pt(10.5)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(43, 43, 43)
    section_style.paragraph_format.space_before = Pt(7)
    section_style.paragraph_format.space_after = Pt(3)

    role_style = ensure_style(document, "Role Header", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    role_style.font.name = "Calibri"
    role_style.font.size = Pt(11)
    role_style.font.bold = True
    role_style.paragraph_format.space_before = Pt(6)
    role_style.paragraph_format.space_after = Pt(0)

    company_style = ensure_style(document, "Company Line", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    company_style.font.name = "Calibri"
    company_style.font.size = Pt(10)
    company_style.font.italic = True
    company_style.font.color.rgb = RGBColor(90, 90, 90)
    company_style.paragraph_format.space_after = Pt(3)

    body_style = ensure_style(document, "Body Tight", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    body_style.font.name = "Calibri"
    body_style.font.size = Pt(10)
    body_style.paragraph_format.space_after = Pt(2)
    body_style.paragraph_format.line_spacing = 1.08

    subhead_style = ensure_style(document, "Subhead", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    subhead_style.font.name = "Calibri"
    subhead_style.font.size = Pt(10)
    subhead_style.font.bold = True
    subhead_style.font.color.rgb = RGBColor(31, 54, 95)
    subhead_style.paragraph_format.space_before = Pt(4)
    subhead_style.paragraph_format.space_after = Pt(1)

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(10)
    bullet_style.paragraph_format.space_after = Pt(1)
    bullet_style.paragraph_format.line_spacing = 1.05
    bullet_style.paragraph_format.left_indent = Inches(0.22)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.16)


def add_section_heading(document, text):
    paragraph = document.add_paragraph(text.upper(), style="Section Heading")
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_border(paragraph)


def add_role(document, title, dates, company):
    paragraph = document.add_paragraph(style="Role Header")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run(title)
    paragraph.add_run("\t")
    paragraph.add_run(dates)
    company_paragraph = document.add_paragraph(company, style="Company Line")
    company_paragraph.paragraph_format.keep_with_next = True


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_subhead(document, text):
    paragraph = document.add_paragraph(text, style="Subhead")
    paragraph.paragraph_format.keep_with_next = True


def build_document():
    document = Document()
    set_cell_margins(document.sections[0])
    configure_styles(document)

    document.add_paragraph("Matthew Snoby", style="Resume Name")
    document.add_paragraph(
        "Technical Leader - Cloud Platforms, Secure Delivery, and AI-Enabled Operations",
        style="Resume Title",
    )
    contact = document.add_paragraph(style="Resume Contact")
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact.add_run(
        "Roswell, GA | matt.snoby@icloud.com | github.com/snoby | "
        "linkedin.com/in/mattsnoby | resume.mattsnoby.com"
    )

    add_section_heading(document, "Summary")
    document.add_paragraph(
        "Hands-on technical leader with 25+ years spanning cloud platforms, secure "
        "developer infrastructure, observability, AI-assisted operations, and embedded "
        "Linux. Known for modernizing production systems with minimal disruption, building "
        "secure delivery foundations for regulated environments, and staying close to the "
        "implementation details from architecture through day-two operations.",
        style="Body Tight",
    )

    add_section_heading(document, "Skills")
    skills = [
        "AI & Automation: Agentic agent design and implementation, AI-native development, spec-driven development, RAG, MCP",
        "Cloud & Platforms: Cloud-native architecture, AWS (EC2, Kafka, Elasticsearch, S3, Route53, IAM, Multi-AZ), OpenStack, SaaS/PaaS platforms, FedRAMP design and support",
        "DevSecOps & Secure Delivery: Terraform, Ansible, Drone CI, GitHub Actions, Jenkins, CircleCI, Spinnaker, ArgoCD, HashiCorp Vault, hardened Linux images",
        "Containers & Observability: Kubernetes, KOPS, Docker, Helm, Splunk, Grafana, Prometheus, Loki, Vector, OTEL, PagerDuty",
        "Engineering & Data: Python, JavaScript, Go, Java, C/C++, PostgreSQL, MongoDB, Oracle, Qdrant, SQLite FTS5, DRM/conditional access, Linux drivers",
    ]
    for skill in skills:
        document.add_paragraph(skill, style="Body Tight")

    add_section_heading(document, "Highlights")
    add_bullets(
        document,
        [
            "Reduced a FedRAMP-secure Kubernetes image workflow from a 20+ step GitOps, JIRA, and MFA process to a single confirmation step.",
            "Designed and sustained Drone CI/CD, Terraform Enterprise, and Vault-backed delivery patterns in production for more than five years without major redesign.",
            "Led Ubuntu 20.04-to-24.04 hardening, DUO SSH adoption, and hardened image rollouts across AWS and OpenStack/3AZ environments while maintaining compliance and platform stability.",
            "Closed out more than 70 deliverables across tasks, epics, and bugs in a single year, supporting modernization, security compliance, and platform readiness workstreams.",
            "Owned production observability and Linux platform work for Webex teams, including Splunk operations, hardened image rollouts, and migrations across AWS and OpenStack environments.",
            "Built and operate independent production systems spanning HA PostgreSQL, Prometheus and Loki observability, blockchain infrastructure, and private AI retrieval platforms.",
        ],
    )

    add_section_heading(document, "Professional Experience")
    add_role(document, "Cloud Engineering Technical Leader", "Sep 2010 - Present", "Cisco / Webex - Roswell, GA (Remote)")
    document.add_paragraph(
        "Lead platform engineering work across AI-assisted infrastructure, secure build "
        "systems, observability, developer platforms, DRM, and Linux-based production services.",
        style="Body Tight",
    )
    add_subhead(document, "Platform Automation, Secure Delivery, and AI Operations")
    add_bullets(
        document,
        [
            "Built automations for a FedRAMP-secure Kubernetes image update workflow, compressing a 20+ step GitOps, JIRA, and MFA-heavy process into a single confirmation step and reusing the same pattern for runbook-guided on-call triage.",
            "Designed and deployed Drone CI/CD and Terraform Enterprise for Webex Logging Metrics, integrated build and deployment secrets with HashiCorp Vault, and kept the architecture stable in production for more than five years.",
            "Designed Vault hierarchies, roles, and policies for CI, development, staging, and production environments; advised teams on GitHub Actions secrets management, image hardening, deployment flows, and day-two operations.",
            "Partnered with a GitHub Actions implementation team that had not previously operated a CI/CD platform, bringing practical production experience from Drone, Terraform Enterprise, Vault integration, build image hardening, and Webex deployment workflows to shape a more realistic rollout path.",
            "Translated lessons from running production Drone CI/CD into GitHub Actions design discussions, helping the team account for secrets handling, build images, deployment workflows, reliability concerns, and day-two operational support.",
            "Built AI-assisted workflows for on-call diagnosis and auto-remediation experiments, combining runbooks and operational context to reduce manual troubleshooting effort.",
        ],
    )
    add_subhead(document, "Observability, Linux, and Platform Operations")
    add_bullets(
        document,
        [
            "Owned a production Splunk platform serving approximately 300 Webex users; led Ubuntu 20.04-to-24.04 hardening, Duo SSH enablement, secure build evidence, and image migrations across AWS and OpenStack/3AZ environments.",
            "Learned and stabilized the LMA logging stack, moved it to hardened images in AWS and 3AZ/OpenStack, resolved authentication and bootstrap issues, and adapted the software stack for newer Ubuntu libraries and operating models.",
            "Automated in-place platform upgrades and validation steps for complex logging environments, including handling reboots, OpenSearch node loss scenarios, and dependency verification across telegraf, AWS CLI, Python, and Ansible.",
            "Moved internal workloads to Kubernetes on AWS before EKS was available, automated multi-region clusters with Ansible and KOPS, and designed a dedicated platform for Webex bot workloads using Kube2IAM and Vault isolation.",
            "Refactored Terraform workspaces, modernized Ansible and image pipelines, and removed friction for internal developer teams by treating platform consumers as customers and proactively fixing tooling bottlenecks.",
            "Delivered more than 70 completed work items in a year across logging, metrics, external platform work, compliance, and infrastructure readiness efforts.",
        ],
    )
    add_subhead(document, "Embedded and Media Platform Foundations")
    add_bullets(
        document,
        [
            "Created custom Linux distribution foundations for Cisco set-top box platforms and wrote Linux kernel and platform drivers for embedded video devices.",
            "Served as technical leader and architect in Cisco's RDK effort, helping shape platform direction for next-generation set-top box software.",
            "Served as technical leader for Cisco Android-based set-top box efforts, guiding architecture and integration decisions across the platform stack.",
            "Represented Cisco in the Linaro LHG open-source collaboration group, including presenting on Chromium and CDM architecture in Hong Kong.",
            "Architected the GStreamer encrypted media pipeline and DRM plugin model used for protected playback workflows.",
            "Led engineering for Intertrust DRM and Nagravision conditional access integrations across Cisco video platforms.",
        ],
    )

    add_role(document, "Principal Engineer", "Sep 2009 - Sep 2010", "Vtilt")
    add_bullets(
        document,
        [
            "Engaged as Cisco's third-wave DRM expert and Linux distribution architect for embedded set-top platforms.",
        ],
    )

    add_role(document, "Associate Staff Engineer", "Aug 2003 - Sep 2009", "Scientific Atlanta / Cisco Systems")
    add_bullets(
        document,
        [
            "Led hardware bring-up and board support package development for dual-processor SPARC ASIC set-top box platforms, including memory map layout, makery design and implementation, and porting existing code to dual-processor compatibility.",
            "Served as team lead for new dual-processor driver designs and code reviews, and carried driver platform responsibility for dual-processor set-top box software delivered to Cablevision in New York.",
            "Designed and implemented DMA engine support across multiple CPUs and wrote Linux kernel character, network, and proprietary next-generation Scientific Atlanta drivers.",
            "Acted as DOCSIS team associate architect and mentor for new hires, including bring-up of a new DOCSIS co-processor and design of the ST40 and ST20 co-processor bootloader.",
            "Built proof-of-concept systems including DOCSIS over USB Ethernet dongle support, dual-processor to single-processor DOCSIS platform conversions, USB dongle bridging, and 802.11b Wi-Fi access point enablement on legacy deployed products.",
            "Supported OCAP COX Cable field trials and demoed advanced proof-of-concept platform capabilities to customers including Cablevision.",
        ],
    )

    add_role(document, "Digital Television Electrical Engineer", "Jan 2001 - Aug 2003", "Livewire")
    add_bullets(
        document,
        [
            "Designed embedded set-top box drivers and resident applications, integrated Nagravision conditional access, and parsed MPEG/DVB service information across TCP/IP, MIPS, ARM, DSP, and ST5518 platforms.",
        ],
    )

    add_role(document, "Software Engineer", "May 2000 - Jan 2001", "Barco")
    add_bullets(
        document,
        [
            "Developed 8051 embedded C software for broadcast systems that digitized and transported audio and video over fiber, and redesigned LCD interface requirements.",
        ],
    )

    add_section_heading(document, "Independent Projects")
    add_subhead(document, "Vipor Mining Pool - vipor.net")
    add_bullets(
        document,
        [
            "Built and operate production mining-pool infrastructure spanning blockchain nodes, stratum services, HAProxy and Nginx routing, Cloudflare and DNS, centralized Vector-to-Loki logging, Prometheus and Grafana monitoring, Alertmanager, and PagerDuty.",
            "Operate PostgreSQL 18 with Patroni primary and replica high availability; executed a no-downtime migration from RDS to EC2 to colocation using logical replication while supporting revenue-critical operations.",
        ],
    )
    add_subhead(document, "Cuckoo Cycle Algorithm R&D")
    add_bullets(
        document,
        [
            "Wrote the world-fastest CPU miner for the Cuckoo mining algorithm by combining CPU affinity and NUMA optimization, compiler tuning, low-level profiling, and benchmark-driven improvements on Ryzen 7950X3D and EPYC systems.",
        ],
    )
    add_subhead(document, "Local AI Systems Lab")
    add_bullets(
        document,
        [
            "Run private Qwen 3.x models with llama.cpp, GGUF quantization, speculative decoding, long-context tuning, and KV-cache optimization for coding and operations workflows.",
            "Designed, built, deployed, and operate DIP, a personal document ingestion and retrieval platform that converts mixed office documents into searchable, cited engineering knowledge through FastAPI.",
            "Authored the DIP architecture covering parallel conversion, SHA256 deduplication, sentence-aware chunking, batched embeddings, Qdrant vector storage, and SQLite FTS5 document state and keyword indexing.",
            "Implemented hybrid retrieval by combining Qdrant vector search and SQLite FTS5 results with reciprocal-rank fusion, chunk deduplication, optional cross-encoder reranking, and OpenAI-compatible cited answer generation.",
            "Operationalized DIP with persistent storage, health and metrics endpoints, request telemetry, failure retry workflows, and systemd-driven incremental ingestion every 30 minutes.",
        ],
    )

    add_section_heading(document, "Education")
    document.add_paragraph("Bachelor of Science, Computer Engineering", style="Body Tight")
    document.add_paragraph("Southern Polytechnic State University - Marietta, Georgia", style="Body Tight")

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
