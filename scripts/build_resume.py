from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "Matthew-Snoby-Resume.docx"

INK = RGBColor(20, 32, 48)
BLUE = RGBColor(15, 92, 132)
CYAN = RGBColor(0, 126, 167)
MUTED = RGBColor(77, 91, 108)
LIGHT = "DCEAF0"
FONT = "Aptos"


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=9.2, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url, color=CYAN, underline=True, size=9.2, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_node)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), str(color))
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def paragraph_border_bottom(paragraph, color="0F5C84", size="12", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run(r, size=10.2, bold=True, color=BLUE)
    paragraph_border_bottom(p, size="8", space="3")
    return p


def add_bullet(doc, text, size=9.0, after=1.7):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run(r, size=size)
    return p


def add_role(doc, company, title, dates, summary=None, bullets=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{company} | {title}")
    set_run(r, size=10, bold=True, color=INK)
    r = p.add_run(f"  |  {dates}")
    set_run(r, size=9, bold=True, color=MUTED)
    if summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = bool(bullets)
        r = p.add_run(summary)
        set_run(r, size=9.0, italic=True, color=MUTED)
    for bullet in bullets or []:
        add_bullet(doc, bullet)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.line_spacing = 1.03
    bullet.paragraph_format.space_after = Pt(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Matthew Snoby  |  Cloud Infrastructure & Platform Engineering Technical Leader")
    set_run(r, size=7.5, color=MUTED)


def add_header(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("MATTHEW SNOBY")
    set_run(r, size=24, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Cloud Infrastructure & Platform Engineering Technical Leader")
    set_run(r, size=12.5, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run("Roswell / Atlanta, GA  |  ")
    add_hyperlink(p, "matt.snoby@gmail.com", "mailto:matt.snoby@gmail.com", size=8.8)
    p.add_run("  |  ")
    add_hyperlink(p, "github.com/snoby", "https://github.com/snoby", size=8.8)
    paragraph_border_bottom(p, size="14", space="6")


def build():
    doc = Document()
    configure_document(doc)
    add_header(doc)

    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(
        "Technical leader with 25+ years spanning cloud platforms, secure developer infrastructure, "
        "Kubernetes, CI/CD, observability, AI-assisted operations, and embedded Linux. Builds durable systems, "
        "modernizes critical platforms without service disruption, and remains hands-on from architecture through production support."
    )
    set_run(r, size=9.3)

    add_section_heading(doc, "Selected Impact")
    impact = [
        ("AI & Knowledge Systems", "Built Codex-assisted FedRAMP workflows and on-call triage; designed and operate DIP as a personal FastAPI document ingestion and hybrid retrieval platform using Qdrant, SQLite FTS5, reranking, and cited answers."),
        ("Reliable Developer Platforms", "Designed and operated Drone CI/CD, Terraform Enterprise, Vault-backed secrets, hardened build images, and GitOps workflows supporting mission-critical Webex engineering."),
        ("Production Operations", "Owned a Splunk platform serving about 300 Webex users and run mining infrastructure with Prometheus, Grafana, Loki, PagerDuty, and PostgreSQL 18 Patroni high availability."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.1)
    for label, detail in impact:
        cells = table.add_row().cells
        cells[0].width = Inches(1.45)
        cells[1].width = Inches(5.1)
        shade_cell(cells[0], LIGHT)
        for cell in cells:
            set_cell_margins(cell)
        p = cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run(r, size=8.8, bold=True, color=BLUE)
        p = cells[1].paragraphs[0]
        r = p.add_run(detail)
        set_run(r, size=8.8)

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc,
        "Cisco / Webex",
        "Cloud Engineering Technical Leader",
        "Sep 2010 - Present",
        "Technical leadership across AI-assisted infrastructure, cloud platforms, secure build systems, observability, DRM, and embedded architecture.",
        [
            "Designed, built, deployed, and operate DIP as a personal project, converting mixed office documents into searchable engineering knowledge with parallel conversion, SHA256 deduplication, sentence-aware chunking, batched embeddings, Qdrant vectors, SQLite FTS5, hybrid retrieval, reranking, and cited answer generation.",
            "Built Codex automations for a FedRAMP-secure Kubernetes image update workflow, reducing a 20+ step GitOps/JIRA/MFA-heavy process to enabling the automation and confirming MFA; also applied Codex to runbook-guided on-call investigation.",
            "Designed and deployed Drone CI/CD and Terraform Enterprise for Webex Logging Metrics, integrated build and deployment secrets with HashiCorp Vault, and maintained the architecture in production for more than five years without major redesign.",
            "Designed Vault hierarchies, roles, and hand-written policies for CI, development, staging, and production; advised a team implementing GitHub Actions on secrets, images, reliability, deployment flow, and day-two operations.",
            "Owned a production Splunk platform for approximately 300 Webex users and drove Ubuntu 20.04-to-24.04 hardening, DUO SSH, secure build evidence, and image migrations across AWS and OpenStack/3AZ.",
            "Moved internal workloads to Kubernetes on AWS before EKS, automating multi-region clusters with Ansible and KOPS; designed a dedicated business-critical platform for Webex bot workloads with Kube2IAM and Vault isolation.",
            "Earlier Cisco leadership included RDK and Android set-top box architecture, GStreamer DRM pipeline plugins, Intertrust DRM and Nagravision integrations, and open source collaboration with Linaro LHG."
        ],
    )

    add_role(
        doc,
        "Vtilt",
        "Principal Engineer",
        "Sep 2009 - Sep 2010",
        "Cisco third-wave DRM expert and Linux distribution architect.",
    )
    doc.add_page_break()
    add_role(
        doc,
        "Scientific Atlanta / Cisco Systems",
        "Associate Staff Engineer",
        "Aug 2003 - Sep 2009",
        "Led low-level platform work across dual-processor set-top box systems, Linux kernel drivers, DOCSIS, DMA, co-processors, and hardware bring-up.",
        [
            "Owned board support package bring-up for dual-processor SPARC ASIC platforms, led driver design and code reviews, and carried platform delivery responsibility for Cablevision deployments.",
            "Designed multi-CPU DMA functionality, Linux character/network/proprietary drivers, DOCSIS co-processor boot flows, and DOCSIS-over-USB and 802.11b proof-of-concept systems."
        ],
    )

    add_section_heading(doc, "Earlier Experience")
    add_role(
        doc,
        "Livewire",
        "Digital Television Electrical Engineer",
        "Jan 2001 - Aug 2003",
        "Designed embedded set-top box drivers and resident applications; integrated Nagravision conditional access, parsed MPEG/DVB service information, and worked across TCP/IP, Win32, MIPS, ARM, DSP, and ST5518 platforms.",
    )
    add_role(
        doc,
        "Barco",
        "Software Engineer",
        "May 2000 - 2001",
        "Developed 8051 embedded C software for broadcast systems that digitized and transported audio/video over fiber, including requirements and LCD interface redesign.",
    )

    add_section_heading(doc, "Independent Engineering")
    add_role(
        doc,
        "Vipor Mining Pool",
        "Infrastructure, Observability & Database Reliability",
        "Independent Project",
        bullets=[
            "Built and operate production mining-pool infrastructure spanning blockchain nodes, stratum services, HAProxy/Nginx routing, Cloudflare/DNS, centralized Vector.dev-to-Loki logging, Prometheus/Grafana monitoring, Alertmanager, and PagerDuty.",
            "Operate PostgreSQL 18 with Patroni primary/replica high availability; migrated from RDS to EC2 and then colocation without downtime using logical replication while supporting millions of dollars in pool revenue."
        ],
    )
    add_role(
        doc,
        "Cuckoo Mining Algorithm",
        "World-Fastest CPU Miner",
        "Infrastructure R&D",
        bullets=[
            "Wrote the world's fastest CPU miner for the Cuckoo mining algorithm by combining affinity and NUMA optimization, compiler tuning, low-level profiling, and benchmark-driven algorithm improvements on Ryzen and EPYC systems."
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(2)
    add_hyperlink(
        p,
        "LinkedIn article: Cuckoo CPU miner performance work",
        "https://www.linkedin.com/posts/activity-7447046008766656512-iTnI?utm_source=share&utm_medium=member_desktop&rcm=ACoAAAIDC_MB0o2CMCkZMqKomzZaBT4C1oRztvo",
        size=8.7,
    )
    add_role(
        doc,
        "Local AI Systems Lab",
        "Private LLM & Retrieval Infrastructure",
        "AI Infrastructure R&D",
        bullets=[
            "Run Qwen 3.x models through llama.cpp with GGUF, MTP speculative decoding, long-context and KV-cache tuning; evaluate private coding and operations workflows across multi-GPU and single-GPU systems."
        ],
    )

    add_section_heading(doc, "Technical Skills")
    skills = [
        ("Cloud & Platforms", "AWS, OpenStack, Kubernetes/KOPS, Terraform, Ansible, Vault, hardened Linux images"),
        ("Developer Infrastructure", "Drone CI, GitHub Actions, Jenkins, CircleCI, GitOps, secure build systems"),
        ("Observability & Data", "Splunk, Prometheus, Grafana, Loki, Vector.dev, PagerDuty, PostgreSQL, Patroni, Qdrant, SQLite FTS5"),
        ("Systems & AI", "C/C++, Python, Linux kernel drivers, DOCSIS, DRM/conditional access, FastAPI, local LLM inference, GGUF, RAG, reranking"),
    ]
    for label, value in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(f"{label}: ")
        set_run(r, size=9, bold=True, color=BLUE)
        r = p.add_run(value)
        set_run(r, size=9)

    add_section_heading(doc, "Education & Selected Links")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("B.S. Computer Engineering, Southern Polytechnic State University, Marietta, Georgia | 2000")
    set_run(r, size=9.2, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_hyperlink(p, "Vipor.net", "https://vipor.net", size=8.8)
    p.add_run("  |  ")
    add_hyperlink(p, "Chromium/CDM Presentation", "https://youtu.be/dJqCbTfKrMk", size=8.8)

    core = doc.core_properties
    core.title = "Matthew Snoby Resume"
    core.subject = "Cloud Infrastructure and Platform Engineering Technical Leader"
    core.author = "Matthew Snoby"
    core.keywords = "Cloud infrastructure, platform engineering, Kubernetes, CI/CD, AI, observability, embedded Linux"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
